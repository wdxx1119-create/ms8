"""Safe parsers for authorized local documents."""

from __future__ import annotations

import ast
import hashlib
import json
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

from .ocr import IMAGE_TYPES, ocr_image, ocr_pdf

MAX_READ_BYTES = 2_000_000
MAX_FILE_BYTES = 50_000_000
TEXT_TYPES = {".md", ".txt", ".yaml", ".yml"}
CODE_TYPES = {".py", ".js", ".ts", ".tsx", ".jsx", ".go", ".rs", ".java", ".c", ".cpp", ".h", ".hpp"}
SUPPORTED_TYPES = TEXT_TYPES | {".json", ".pdf", ".docx"} | CODE_TYPES | IMAGE_TYPES


@dataclass
class ParsedDocument:
    source_path: str
    file_type: str
    title: str
    content_text: str
    content_hash: str
    metadata: dict[str, Any] = field(default_factory=dict)
    parse_status: str = "parsed"
    error: str = ""


def _hash_text(text: str) -> str:
    return hashlib.sha256(str(text or "").encode("utf-8")).hexdigest()


def _read_limited(path: Path) -> str:
    with path.open("rb") as fh:
        raw = fh.read(MAX_READ_BYTES)
    return raw.decode("utf-8", errors="replace")


def _parse_json(path: Path) -> str:
    payload = json.loads(_read_limited(path))
    return json.dumps(payload, ensure_ascii=False, indent=2)


def _parse_python_docs(path: Path) -> str:
    text = _read_limited(path)
    try:
        tree = ast.parse(text)
    except SyntaxError:
        return "\n".join(line.strip() for line in text.splitlines() if line.strip().startswith("#"))
    parts: list[str] = []
    module_doc = ast.get_docstring(tree)
    if module_doc:
        parts.append(module_doc)
    for node in ast.walk(tree):
        if isinstance(node, (ast.FunctionDef, ast.AsyncFunctionDef, ast.ClassDef)):
            doc = ast.get_docstring(node)
            if doc:
                parts.append(f"{node.name}: {doc}")
    comments = [line.strip("# ").strip() for line in text.splitlines() if line.strip().startswith("#")]
    parts.extend(c for c in comments if c)
    return "\n".join(parts)


def _parse_js_docs(path: Path) -> str:
    text = _read_limited(path)
    block_comments = re.findall(r"/\*+(.*?)\*/", text, flags=re.DOTALL)
    line_comments = re.findall(r"^\s*//\s*(.+)$", text, flags=re.MULTILINE)
    parts = [re.sub(r"^\s*\* ?", "", c, flags=re.MULTILINE).strip() for c in block_comments]
    parts.extend(c.strip() for c in line_comments)
    return "\n".join(p for p in parts if p)


def _parse_generic_code_docs(path: Path) -> str:
    text = _read_limited(path)
    block_comments = re.findall(r"/\*+(.*?)\*/", text, flags=re.DOTALL)
    line_comments = re.findall(r"^\s*(?://|#|--)\s*(.+)$", text, flags=re.MULTILINE)
    parts = [re.sub(r"^\s*\* ?", "", c, flags=re.MULTILINE).strip() for c in block_comments]
    parts.extend(c.strip() for c in line_comments)
    if parts:
        return "\n".join(p for p in parts if p)
    return "\n".join(line.strip() for line in text.splitlines() if line.strip())[:MAX_READ_BYTES]


def _parse_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise RuntimeError("missing_optional_dependency:pypdf") from exc
    reader = PdfReader(str(path))
    texts: list[str] = []
    for page in reader.pages[:50]:
        texts.append(page.extract_text() or "")
    return "\n".join(t for t in texts if t.strip())


def _parse_docx(path: Path) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("missing_optional_dependency:python-docx") from exc
    doc = Document(str(path))
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())


def _ocr_required(path: Path, suffix: str, title: str, reason: str = "ocr_required") -> ParsedDocument:
    return ParsedDocument(
        source_path=str(path),
        file_type=suffix,
        title=title,
        content_text="",
        content_hash="",
        metadata={"size": path.stat().st_size, "mtime": path.stat().st_mtime},
        parse_status="ocr_required",
        error=reason,
    )


def parse_document(path: str | Path) -> ParsedDocument:
    p = Path(path).expanduser().resolve()
    suffix = p.suffix.lower()
    try:
        if suffix not in SUPPORTED_TYPES:
            return ParsedDocument(str(p), suffix, p.name, "", "", parse_status="unsupported", error="unsupported_type")
        if suffix in TEXT_TYPES:
            content = _read_limited(p)
        elif suffix == ".json":
            content = _parse_json(p)
        elif suffix == ".pdf":
            try:
                content = _parse_pdf(p)
            except RuntimeError as exc:
                if "missing_optional_dependency:pypdf" not in str(exc):
                    raise
                try:
                    content = ocr_pdf(p)
                except RuntimeError as ocr_exc:
                    return _ocr_required(p, suffix, p.stem, str(ocr_exc))
            if not content.strip():
                try:
                    content = ocr_pdf(p)
                except RuntimeError as exc:
                    return _ocr_required(p, suffix, p.stem, str(exc))
        elif suffix == ".docx":
            content = _parse_docx(p)
        elif suffix in IMAGE_TYPES:
            try:
                content = ocr_image(p)
            except RuntimeError as exc:
                return _ocr_required(p, suffix, p.stem, str(exc))
        elif suffix == ".py":
            content = _parse_python_docs(p)
        elif suffix in {".js", ".ts"}:
            content = _parse_js_docs(p)
        elif suffix in CODE_TYPES:
            content = _parse_generic_code_docs(p)
        else:
            content = ""
        title = p.stem
        if suffix in {".pdf", *IMAGE_TYPES} and not content.strip():
            return _ocr_required(p, suffix, title)
        return ParsedDocument(
            source_path=str(p),
            file_type=suffix,
            title=title,
            content_text=content,
            content_hash=_hash_text(content),
            metadata={"size": p.stat().st_size, "mtime": p.stat().st_mtime},
            parse_status="parsed" if content.strip() else "empty",
        )
    except (OSError, RuntimeError, UnicodeDecodeError, json.JSONDecodeError, ValueError) as exc:
        return ParsedDocument(str(p), suffix, p.name, "", "", parse_status="error", error=str(exc))
